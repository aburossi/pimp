# app/output_generator.py

import os
from docx import Document as DocxDocument
from docx.shared import Inches

def create_docx(content_data, task_type):
    """Generates a .docx file from the structured data."""
    doc = DocxDocument()
    doc.add_heading(content_data.title, level=1)

    if task_type == "Lektion erstellen":
        doc.add_heading("Learning Objectives", level=2)
        for obj in content_data.learning_objectives:
            doc.add_paragraph(obj, style='List Bullet')
        doc.add_heading("Key Concepts", level=2)
        for concept in content_data.key_concepts:
            doc.add_paragraph(concept, style='List Bullet')
        doc.add_heading("Activities", level=2)
        for activity in content_data.activities:
            doc.add_paragraph(activity, style='List Bullet')

    elif task_type == "Quiz erstellen":
        for i, q in enumerate(content_data.questions, 1):
            doc.add_heading(f"Question {i}: {q.question_text}", level=3)
            for opt in q.options:
                doc.add_paragraph(f"- {opt}")
            p = doc.add_paragraph(f"Correct Answer: {q.correct_answer}")
            p.runs[0].bold = True
            doc.add_paragraph()
    
    temp_dir = "temp"
    if not os.path.exists(temp_dir): os.makedirs(temp_dir)
    file_path = os.path.join(temp_dir, f"{content_data.title.replace(' ', '_')}.docx")
    doc.save(file_path)
    return file_path

def create_html(content_data, task_type):
    """Generates an .html file from the structured data."""
    html = f"<html><head><title>{content_data.title}</title>"
    html += "<style>body { font-family: sans-serif; margin: 2em; } h1, h2, h3 { color: #333; } ul { list-style-type: disc; margin-left: 20px; }</style>"
    html += f"</head><body><h1>{content_data.title}</h1>"

    if task_type == "Lektion erstellen":
        html += "<h2>Learning Objectives</h2><ul>" + "".join(f"<li>{obj}</li>" for obj in content_data.learning_objectives) + "</ul>"
        html += "<h2>Key Concepts</h2><ul>" + "".join(f"<li>{concept}</li>" for concept in content_data.key_concepts) + "</ul>"
        html += "<h2>Activities</h2><ul>" + "".join(f"<li>{activity}</li>" for activity in content_data.activities) + "</ul>"
    elif task_type == "Quiz erstellen":
        for i, q in enumerate(content_data.questions, 1):
            html += f"<h3>Question {i}: {q.question_text}</h3><ul>" + "".join(f"<li>{opt}</li>" for opt in q.options) + "</ul>"
            html += f"<p><b>Correct Answer:</b> {q.correct_answer}</p><hr>"

    html += "</body></html>"
    temp_dir = "temp"
    if not os.path.exists(temp_dir): os.makedirs(temp_dir)
    file_path = os.path.join(temp_dir, f"{content_data.title.replace(' ', '_')}.html")
    with open(file_path, "w", encoding='utf-8') as f:
        f.write(html)
    return file_path