import re
import yaml
from urllib.parse import urlparse, parse_qs, unquote_plus
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Helper Function to create a shaded box for structure ---
def add_shaded_box(document, title, text_lines):
    """
    Creates a table with one cell, a shaded background, and a border 
    to act as a visual container for content.
    """
    table = document.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)

    # Set the background color of the cell to light gray
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:fill'), 'EAEAEA') # Light gray fill
    cell._tc.get_or_add_tcPr().append(shading_elm)

    # Add the title to the cell
    p_title = cell.add_paragraph()
    p_title.add_run(title).bold = True
    
    # Add the content text to the cell
    for line in text_lines:
        # Clean up common markdown from the text
        cleaned_line = re.sub(r'(\*\*|\[\[|\]\])', '', line)
        cleaned_line = re.sub(r'\s*\- ', '', cleaned_line) # Remove list markers
        cell.add_paragraph(cleaned_line.strip())

# --- Helper Function to parse questions from an iFrame URL ---
def add_questions_from_iframe(document, url):
    """
    Parses a URL, extracts question parameters, decodes them, and formats
    them for a printable worksheet.
    """
    parsed_url = urlparse(url)
    query_params = parse_qs(parsed_url.query)
    
    # Find all parameters that start with 'question'
    question_keys = sorted([k for k in query_params.keys() if k.startswith('question')])

    if not question_keys:
        return

    document.add_heading("Fragen zum Beantworten", level=3)
    
    for i, key in enumerate(question_keys, 1):
        # Decode URL-encoded text (e.g., %20 -> space)
        encoded_text = query_params[key][0]
        decoded_text = unquote_plus(encoded_text)
        
        # Clean up markdown formatting (** for bold, * for italic)
        clean_text = re.sub(r'(\*\*|\*)', '', decoded_text)
        
        # Add the question to the document
        p = document.add_paragraph()
        p.add_run(f"{i}. {clean_text}").bold = True
        
        # Add blank lines for a printed answer
        # A table with a bottom border creates clean lines for writing
        table = document.add_table(rows=4, cols=1)
        # Apply a style that only has bottom borders (or create one)
        # For simplicity, we just leave blank paragraphs in the cells
        for row in table.rows:
            # Removing the default paragraph margins to make the lines tighter
            p = row.cells[0].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)


# --- Main Parsing and Generation Function ---
def create_printable_word_doc(md_file_path, docx_file_path):
    """
    Parses a structured Markdown file and generates a clean, printable Word document.
    """
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 1. Separate body from frontmatter, but we will not use the frontmatter.
    try:
        _, _, body_str = content.split('---', 2)
    except ValueError:
        body_str = content

    document = Document()

    # 2. Process Markdown Body
    lines = body_str.strip().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue

        # --- Map Markdown/Custom elements to DOCX elements ---

        # Heading Levels
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('# ').strip()
            # Don't create a title bigger than Heading 1
            document.add_heading(title, level=min(level, 4))
            i += 1

        # Custom Callouts -> Shaded Box
        elif match := re.match(r'>\[!(\w+)\](.*)', line):
            title = match.group(2).strip()
            # Gather subsequent lines belonging to the same callout block
            content_lines = []
            i += 1
            while i < len(lines) and lines[i].strip().startswith('>>'):
                content_lines.append(lines[i].strip().lstrip('> '))
                i += 1
            
            # If no content lines, maybe the content is on the same line (older format)
            if not content_lines and len(line.split(']')) > 1:
                 # Check for content on the initial callout line itself
                 potential_content = line.split(']', 1)[1].strip()
                 if potential_content:
                     content_lines.append(potential_content)

            add_shaded_box(document, title, content_lines)

        # iFrames -> Extracted Questions
        elif '<iframe src="' in line:
            if src_match := re.search(r'src="([^"]+)"', line):
                url = src_match.group(1)
                add_questions_from_iframe(document, url)
            i += 1

        # Audio Files -> Simple Link
        elif '<audio controls>' in line:
            if src_match := re.search(r'<source src="([^"]+)"', line):
                url = src_match.group(1)
                document.add_paragraph(f"Zugehöriger Radiobeitrag: {url}")
            i += 1
            
        # Any other text
        else:
            # Clean up wiki-links and other simple markdown
            clean_line = re.sub(r'\[\[(.*?)\]\]', r'\1', line)
            if clean_line.strip():
                document.add_paragraph(clean_line)
            i += 1
            
    # Add a final paragraph for spacing
    document.add_paragraph()
    
    # 3. Save the document
    document.save(docx_file_path)
    print(f"Successfully created printable Word document at: {docx_file_path}")


# --- Execution ---
if __name__ == '__main__':
    markdown_file = 'your_markdown_file.md'
    word_file = 'printable_worksheet.docx'
    
    # Re-create the markdown file for demonstration purposes
    markdown_content = """
---
topic: ["Politik", "Recht", "Gewaltentrennung", "Politik und Umfeld", "Migration"]
chapter: []
type: "news"
source: "ABUnews"
summary: "**Worum es geht**: In diesem Radiobeitrag geht es um die **Proteste** in Los Angeles gegen die Einwanderungsbehörde **ICE**, die eskalierten, nachdem US-Präsident Trump gegen den Willen der kalifornischen Behörden die **Nationalgarde** einsetzte."
---
# ABUnews - Unruhen in den USA
>[!info] Worum geht es?
>> Das Thema dieses Radiobeitrags sind die **heftigen Proteste** in der amerikanischen Stadt Los Angeles. Es geht um einen Konflikt zwischen dem US-Präsidenten und den **Behörden von Kalifornien**, der zeigt, wie unterschiedlich mit Menschen ohne gültige Papiere umgegangen wird.
>
>>[!success] Lernziele
>>- Sie können die Hauptgründe für die Proteste in Los Angeles und die Anliegen der Demonstrierenden erklären.
>>- Sie können die Rolle der Nationalgarde in den USA erläutern und bewerten, warum ihr Einsatz in diesem Fall als aussergewöhnlich gilt.
>
>#### Schlüsselbegriffe
>[[Gewaltentrennung]], [[Politische Rechte]], [[Staatsformen]]
>
>[!question] Wählen Sie **eine Frage** und **begründen Sie** Ihre Antwort 
><iframe src="https://allgemeinbildung.github.io/textbox/answers.html?assignmentId=Unruhen%20in%20den%20USA&subIds=A.%20Einstiegsfragen%20Unruhen%20in%20den%20USA&question1=**Haben%20Sie%20schon%20einmal%20an%20einer%20Demonstration%20teilgenommen%20oder%20darüber%20nachgedacht?%20Was%20sind%20Ihrer%20Meinung%20nach%20gute%20Gründe,%20um%20auf%20die%20Strasse%20zu%20gehen%20und%20zu%20protestieren?**%20*Wie%20unterscheidet%20sich%20friedlicher%20Protest%20von%20Randalen,%20und%20wo%20ziehen%20Sie%20persönlich%20die%20Grenze?*&question2=**In%20der%20Schweiz%20gibt%20es%20klare%20Regeln,%20wer%20die%20Armee%20oder%20die%20Polizei%20einsetzen%20darf.%20Was%20denken%20Sie,%20was%20passieren%20würde,%20wenn%20sich%20die%20Regierung%20in%20Bern%20über%20den%20Willen%20eines%20Kantons%20hinwegsetzen%20würde,%20um%20zum%20Beispiel%20die%20Armee%20aufzubieten?**&question3=**Menschen%20ohne%20gültige%20Papiere%20('Sans-Papiers')%20leben%20auch%20in%20der%20Schweiz.%20Welche%20Herausforderungen%20und%20Ängste%20haben%20diese%20Menschen%20Ihrer%20Meinung%20nach%20im%20Alltag?**" style="border:0px #ffffff none;" name="myiFrame" scrolling="yes" frameborder="1" marginheight="0px" marginwidth="0px" height="450px" width="100%" allowfullscreen></iframe>
"""
    with open(markdown_file, 'w', encoding='utf-8') as f:
        f.write(markdown_content)

    create_printable_word_doc(markdown_file, word_file)